from pathlib import Path
from docx import Document


class TemplateService:
    STAGE_CODES = {
        "Draft": "v0.1",
        "Reviewed": "v0.2",
        "Revised": "v0.3",
    }

    def parse_template(self, template_name: str) -> dict:
        return {
            "template": template_name,
            "headings": [
                "1. Purpose",
                "2. Intended Use",
                "3. Performance Claims",
                "4. Risk Summary",
                "5. Evidence Table",
            ],
            "placeholders": ["{{PROJECT_NAME}}", "{{DOCUMENT_TYPE}}"],
            "table_mapping": {"Document Type": "{{DOCUMENT_TYPE}}", "Status": "{{STATUS}}"},
            "editable_regions": ["BODY", "TABLE_EVIDENCE"],
        }

    def build_versioned_output_path(self, output_root: str, document_type: str, stage: str) -> tuple[str, str]:
        safe_doc_type = document_type.replace(" ", "_")
        version = self.STAGE_CODES.get(stage, "v0.1")
        filename = f"{safe_doc_type}_{stage}_{version}.docx"
        root = Path(output_root)
        if root.suffix.lower() == ".docx":
            root = root.parent
        root.mkdir(parents=True, exist_ok=True)
        return str(root / filename), version

    def generate_docx(
        self,
        template_name: str,
        output_root: str,
        document_type: str,
        stage: str,
        sections: list[dict],
        placeholders: dict,
        change_summary: list[str],
        revision_context: dict | None = None,
        template_path_override: str | None = None,
    ) -> tuple[str, str]:
        output_path, version = self.build_versioned_output_path(output_root, document_type, stage)
        template_path = Path(template_path_override) if template_path_override else Path("backend/data/templates") / template_name

        # Read template and only save a new output file; source template is never modified.
        document = Document(template_path) if template_path.exists() else Document()
        placeholder_values = {"{{DOCUMENT_TYPE}}": document_type, "{{STATUS}}": stage}
        placeholder_values.update(placeholders)

        for paragraph in document.paragraphs:
            for key, value in placeholder_values.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in placeholder_values.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, value)

        document.add_heading(f"{document_type} - {stage}", level=1)
        for section in sections:
            self._insert_section_content(document, section)

        self._append_revision_summary(document, change_summary, revision_context or {})

        document.save(output_path)
        return output_path, version

    def _insert_section_content(self, document: Document, section: dict) -> None:
        title = section["section_title"]
        heading = self._find_heading(document, title)
        if not heading:
            heading = document.add_heading(title, level=2)

        insertion_point = heading
        insertion_point = self._insert_paragraph_after(insertion_point, section["generated_text"], "Normal")

        refs = ", ".join(section.get("evidence_refs", [])) if section.get("evidence_refs") else "No linked evidence"
        insertion_point = self._insert_paragraph_after(insertion_point, f"Evidence links: {refs}", "Intense Quote")

        evidence_metadata = section.get("evidence_metadata", [])
        if evidence_metadata:
            items = [f"{item.get('file_id', 'unknown')} (score={item.get('score', 'n/a')})" for item in evidence_metadata]
            insertion_point = self._insert_paragraph_after(insertion_point, f"Evidence metadata: {', '.join(items)}", "Intense Quote")

        gaps = section.get("unresolved_gaps", [])
        if gaps:
            self._insert_paragraph_after(insertion_point, "Missing evidence notes: " + "; ".join(gaps), "Intense Quote")

    def _append_revision_summary(self, document: Document, change_summary: list[str], revision_context: dict) -> None:
        document.add_heading("Revision Summary", level=2)
        for line in change_summary:
            document.add_paragraph(line, style="List Bullet")

        addressed = revision_context.get("addressed_findings", [])
        remaining = revision_context.get("remaining_findings", [])
        rationale = revision_context.get("change_rationale", [])

        document.add_paragraph("Addressed findings", style="Heading 3")
        if addressed:
            for item in addressed:
                document.add_paragraph(f"#{item['id']}: {item['category']} - {item['issue_summary']}", style="List Bullet")
        else:
            document.add_paragraph("No findings addressed in this cycle.")

        document.add_paragraph("Remaining findings", style="Heading 3")
        if remaining:
            for item in remaining:
                document.add_paragraph(f"#{item['id']}: {item['category']} - {item['issue_summary']}", style="List Bullet")
        else:
            document.add_paragraph("No open findings remain.")

        document.add_paragraph("Change rationale", style="Heading 3")
        if rationale:
            for line in rationale:
                document.add_paragraph(line, style="List Bullet")
        else:
            document.add_paragraph("No additional rationale captured.")

    def _find_heading(self, document: Document, section_title: str):
        target = section_title.lower().strip()
        for paragraph in document.paragraphs:
            style_name = paragraph.style.name.lower() if paragraph.style else ""
            if "heading" not in style_name:
                continue
            raw = paragraph.text.strip().lower()
            normalized = raw.split(".", 1)[-1].strip() if "." in raw[:3] else raw
            if raw == target or normalized == target:
                return paragraph
        return None

    def _insert_paragraph_after(self, paragraph, text: str, style: str):
        new_p = paragraph.insert_paragraph_before(text)
        paragraph._p.addnext(new_p._p)
        new_p.style = style
        return new_p
